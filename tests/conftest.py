from __future__ import annotations

import io
from pathlib import Path

import pytest


@pytest.fixture
def kb_dir(tmp_path: Path) -> Path:
    (tmp_path / "deploy-guide.md").write_text(
        "# Deployment Guide\n\nTo deploy to staging run the Jenkins job deploy-staging.\n"
        "Production needs a change ticket.\n",
        encoding="utf-8",
    )
    (tmp_path / "ticketing.md").write_text(
        "# Ticketing\n\nCreate a ticket in ServiceNow. Choose Incident or Service Request.\n",
        encoding="utf-8",
    )
    (tmp_path / "notes.html").write_text(
        "<html><head><style>p{}</style><script>var x=1;</script></head>"
        "<body><h1>VPN</h1><p>Raise a service request for VPN access.</p></body></html>",
        encoding="utf-8",
    )
    (tmp_path / "ignored.bin").write_bytes(b"\x00\x01")
    sub = tmp_path / "hr"
    sub.mkdir()
    import docx

    document = docx.Document()
    document.add_heading("Leave Policy", level=1)
    document.add_paragraph("Submit leave requests in Workday at least two weeks ahead.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Type"
    table.rows[0].cells[1].text = "Days"
    buffer = io.BytesIO()
    document.save(buffer)
    (sub / "leave-policy.docx").write_bytes(buffer.getvalue())
    return tmp_path
